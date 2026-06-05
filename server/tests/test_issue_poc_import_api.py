from io import BytesIO

from docx import Document
from flask import Flask
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from models import IssueBug, IssueBugComment, IssuePocProgress, IssuePocProject, IssueProduct, IssueProductMapping
from utils.db import Base
from utils.issue_poc_import import parse_weekly_report


def _json_body(response):
    if isinstance(response, tuple):
        return response[0].get_json()
    return response.get_json()


def _docx_bytes():
    document = Document()
    table = document.add_table(rows=3, cols=11)
    headers = ["项目编码", "客户", "网安产品", "版本", "销售支撑", "本周支撑内容", "风险描述", "分类", "BugID", "闭环方", "状态"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    values = ["POC-1", "客户A", "周报产品A", "V1.0", "销售张", "本周推进", "存在风险", "产品问题", "BUG-1", "研发", "进行中"]
    for index, value in enumerate(values):
        table.cell(1, index).text = value
    values = ["", "客户B", "周报产品A", "V1.1", "销售李", "手工项目", "", "", "", "", "跟进中"]
    for index, value in enumerate(values):
        table.cell(2, index).text = value

    quality = document.add_table(rows=2, cols=3)
    quality.cell(0, 0).text = "BugID"
    quality.cell(0, 1).text = "质量备注"
    quality.cell(0, 2).text = "其它"
    quality.cell(1, 0).text = "BUG-1"
    quality.cell(1, 1).text = "需要补充回归验证"
    quality.cell(1, 2).text = ""

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _setup_session(monkeypatch):
    from blueprints import issue_poc

    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(engine)
    TestingSession = sessionmaker(bind=engine)
    monkeypatch.setattr(issue_poc, "SessionLocal", TestingSession)
    return issue_poc, TestingSession


def test_parse_weekly_report_extracts_projects_and_quality_comments():
    projects, comments, errors = parse_weekly_report(_docx_bytes())

    assert errors == []
    assert len(projects) == 2
    assert projects[0]["project_code"] == "POC-1"
    assert projects[0]["product_raw"] == "周报产品A"
    assert projects[0]["has_risk"] == 1
    assert projects[1]["project_code"] == ""
    assert comments == [{"bug_id": "BUG-1", "comment_text": "需要补充回归验证"}]


def test_import_poc_creates_updates_progress_and_bug_comments(monkeypatch):
    issue_poc, TestingSession = _setup_session(monkeypatch)

    db = TestingSession()
    product = IssueProduct(name="RayScan", description="")
    db.add(product)
    db.commit()
    db.add(IssueProductMapping(product_id=product.id, source_type="weekly", source_name="周报产品A"))
    db.add(IssueBug(bug_id="BUG-1", product_id=product.id, title="登录失败"))
    db.add(
        IssuePocProject(
            project_code="POC-1",
            customer_name="旧客户",
            product_id=product.id,
            version="V0.9",
            weekly_content="上周内容",
            current_status="旧状态",
            source_report="上周周报.docx",
            root_cause="保留根因",
            next_step="保留下步",
        )
    )
    db.commit()
    db.close()

    app = Flask(__name__)
    with app.test_request_context(
        method="POST",
        data={"file": (BytesIO(_docx_bytes()), "本周周报.docx")},
        content_type="multipart/form-data",
    ):
        body = _json_body(issue_poc.import_poc.__wrapped__())

    assert body["code"] == 200
    assert body["data"]["new_count"] == 1
    assert body["data"]["update_count"] == 1
    assert body["data"]["comment_count"] == 1

    db = TestingSession()
    updated = db.query(IssuePocProject).filter(IssuePocProject.project_code == "POC-1").first()
    manual = db.query(IssuePocProject).filter(IssuePocProject.project_code.like("MANUAL-%")).first()
    progress = db.query(IssuePocProgress).filter(IssuePocProgress.project_id == updated.id).first()
    comment = db.query(IssueBugComment).first()

    assert updated.customer_name == "客户A"
    assert updated.weekly_content == "本周推进"
    assert updated.root_cause == "保留根因"
    assert updated.next_step == "保留下步"
    assert manual.customer_name == "客户B"
    assert progress.description == "上周内容"
    assert progress.status == "旧状态"
    assert comment.content == "[周报提取] 需要补充回归验证"
    assert comment.source == "周报提取"
    db.close()


def test_import_poc_route_is_registered(monkeypatch):
    import app as app_module

    monkeypatch.setattr(app_module, "init_db", lambda config: None)
    monkeypatch.setattr(app_module, "init_redis", lambda config: None)

    flask_app = app_module.create_app()
    rules = {rule.rule for rule in flask_app.url_map.iter_rules()}

    assert "/api/issue/poc/import" in rules
